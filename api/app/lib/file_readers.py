"""
ファイル読み込みモジュール
Word、Excel、PDF、テキストファイルからのテキスト抽出
Gemini OCR フォールバックなし / Streamlit 依存なし
"""

import io
import sys
from typing import BinaryIO

# サポートするファイル拡張子
SUPPORTED_EXTENSIONS = {
    '.docx': 'Word文書',
    '.xlsx': 'Excelファイル',
    '.pdf': 'PDFファイル',
    '.txt': 'テキストファイル',
    '.jpg': '画像（JPEG）',
    '.jpeg': '画像（JPEG）',
    '.png': '画像（PNG）',
    '.webp': '画像（WebP）',
}

# 画像拡張子のセット
_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp'}


def _log(message: str, level: str = "INFO") -> None:
    sys.stderr.write(f"[FileReaders:{level}] {message}\n")
    sys.stderr.flush()


def get_supported_extensions() -> dict:
    """サポートするファイル拡張子の辞書を返す"""
    return SUPPORTED_EXTENSIONS


def read_docx(file: BinaryIO) -> str:
    """
    Word文書(.docx)からテキストを抽出

    Args:
        file: バイナリファイルオブジェクト

    Returns:
        抽出されたテキスト

    Raises:
        ImportError: python-docx が未インストールの場合
        ValueError: ファイル読み込み失敗の場合
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docxがインストールされていません。`uv add python-docx`を実行してください。"
        )

    try:
        doc = Document(file)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))

        return '\n\n'.join(paragraphs)

    except Exception as e:
        raise ValueError(f"Word文書の読み込みに失敗しました: {e}")


def read_xlsx(file: BinaryIO) -> str:
    """
    Excelファイル(.xlsx)からテキストを抽出

    Args:
        file: バイナリファイルオブジェクト

    Returns:
        抽出されたテキスト

    Raises:
        ImportError: openpyxl が未インストールの場合
        ValueError: ファイル読み込み失敗の場合
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError(
            "openpyxlがインストールされていません。`uv add openpyxl`を実行してください。"
        )

    try:
        wb = load_workbook(file, read_only=True, data_only=True)
        all_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"【シート: {sheet_name}】"]

            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    sheet_text.append(' | '.join(row_values))

            if len(sheet_text) > 1:
                all_text.append('\n'.join(sheet_text))

        wb.close()
        return '\n\n'.join(all_text)

    except Exception as e:
        raise ValueError(f"Excelファイルの読み込みに失敗しました: {e}")


def read_pdf(file: BinaryIO) -> str:
    """
    PDFファイル(.pdf)からテキストを抽出（pdfplumber 使用）

    Args:
        file: バイナリファイルオブジェクト

    Returns:
        抽出されたテキスト

    Raises:
        ImportError: pdfplumber が未インストールの場合
        ValueError: ファイル読み込み失敗の場合
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumberがインストールされていません。`uv add pdfplumber`を実行してください。"
        )

    try:
        all_text = []
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    all_text.append(f"【ページ {i}】\n{page_text}")

                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row and any(cell for cell in row if cell):
                            row_text = ' | '.join(
                                str(cell) if cell else '' for cell in row
                            )
                            all_text.append(row_text)

        return '\n\n'.join(all_text)

    except Exception as e:
        raise ValueError(f"PDFファイルの読み込みに失敗しました: {e}")


def read_txt(file: BinaryIO) -> str:
    """
    テキストファイル(.txt)を読み込み
    エンコーディングを順に試行: utf-8 → shift_jis → cp932 → euc-jp

    Args:
        file: バイナリファイルオブジェクト

    Returns:
        ファイル内容

    Raises:
        ValueError: ファイル読み込み失敗の場合
    """
    try:
        content = file.read()

        for encoding in ['utf-8', 'shift_jis', 'cp932', 'euc-jp']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        # すべて失敗した場合はエラーを無視して読み込み
        return content.decode('utf-8', errors='ignore')

    except Exception as e:
        raise ValueError(f"テキストファイルの読み込みに失敗しました: {e}")


def read_file(file: BinaryIO, filename: str) -> str:
    """
    ファイル拡張子に応じて適切な読み込み関数にディスパッチ

    Args:
        file: バイナリファイルオブジェクト
        filename: ファイル名（拡張子の判定に使用）

    Returns:
        抽出されたテキスト

    Raises:
        ValueError: サポートされていないファイル形式の場合
        ImportError: 必要なライブラリがない場合
    """
    lower_name = filename.lower()

    # ファイルをBytesIOに変換（再読み込み可能にする）
    file_bytes = io.BytesIO(file.read())

    if lower_name.endswith('.docx'):
        return read_docx(file_bytes)
    elif lower_name.endswith('.xlsx'):
        return read_xlsx(file_bytes)
    elif lower_name.endswith('.pdf'):
        return read_pdf(file_bytes)
    elif lower_name.endswith('.txt'):
        return read_txt(file_bytes)
    else:
        ext = '.' + lower_name.rsplit('.', 1)[-1] if '.' in lower_name else ''
        if ext in _IMAGE_EXTENSIONS:
            raise ValueError(
                f"画像ファイル ({ext}) の読み込みは現在サポートされていません。"
                "テキスト、Word、Excel、PDF形式のファイルを使用してください。"
            )
        supported = ', '.join(SUPPORTED_EXTENSIONS.keys())
        raise ValueError(
            f"サポートされていないファイル形式です: {ext or '不明'}\n対応形式: {supported}"
        )
