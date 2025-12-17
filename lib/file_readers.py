"""
親亡き後支援データベース - ファイル読み込みモジュール
Word、Excel、PDF、テキストファイルからのテキスト抽出
"""

import io
from typing import BinaryIO

# サポートするファイル拡張子
SUPPORTED_EXTENSIONS = {
    '.docx': 'Word文書',
    '.xlsx': 'Excelファイル', 
    '.pdf': 'PDFファイル',
    '.txt': 'テキストファイル',
}


def get_supported_extensions() -> dict:
    """サポートするファイル拡張子を取得"""
    return SUPPORTED_EXTENSIONS


def read_docx(file: BinaryIO) -> str:
    """
    Word文書(.docx)からテキストを抽出
    
    Args:
        file: アップロードされたファイルオブジェクト
        
    Returns:
        抽出されたテキスト
    """
    try:
        from docx import Document
        
        doc = Document(file)
        
        # 段落を抽出
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # テーブルからも抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))
        
        return '\n\n'.join(paragraphs)
        
    except ImportError:
        raise ImportError("python-docxがインストールされていません。`uv add python-docx`を実行してください。")
    except Exception as e:
        raise ValueError(f"Word文書の読み込みに失敗しました: {e}")


def read_xlsx(file: BinaryIO) -> str:
    """
    Excelファイル(.xlsx)からテキストを抽出
    
    Args:
        file: アップロードされたファイルオブジェクト
        
    Returns:
        抽出されたテキスト
    """
    try:
        from openpyxl import load_workbook
        
        wb = load_workbook(file, read_only=True, data_only=True)
        
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"【シート: {sheet_name}】"]
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    sheet_text.append(' | '.join(row_values))
            
            if len(sheet_text) > 1:  # シート名以外に内容がある場合
                all_text.append('\n'.join(sheet_text))
        
        wb.close()
        return '\n\n'.join(all_text)
        
    except ImportError:
        raise ImportError("openpyxlがインストールされていません。`uv add openpyxl`を実行してください。")
    except Exception as e:
        raise ValueError(f"Excelファイルの読み込みに失敗しました: {e}")


def read_pdf(file: BinaryIO) -> str:
    """
    PDFファイル(.pdf)からテキストを抽出
    
    Args:
        file: アップロードされたファイルオブジェクト
        
    Returns:
        抽出されたテキスト
    """
    try:
        import pdfplumber
        
        all_text = []
        
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    all_text.append(f"【ページ {i}】\n{page_text}")
                
                # テーブルがあれば抽出
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row and any(cell for cell in row if cell):
                            row_text = ' | '.join(str(cell) if cell else '' for cell in row)
                            all_text.append(row_text)
        
        return '\n\n'.join(all_text)
        
    except ImportError:
        raise ImportError("pdfplumberがインストールされていません。`uv add pdfplumber`を実行してください。")
    except Exception as e:
        raise ValueError(f"PDFファイルの読み込みに失敗しました: {e}")


def read_txt(file: BinaryIO) -> str:
    """
    テキストファイル(.txt)を読み込み
    
    Args:
        file: アップロードされたファイルオブジェクト
        
    Returns:
        ファイル内容
    """
    try:
        # 複数のエンコーディングを試行
        content = file.read()
        
        for encoding in ['utf-8', 'shift_jis', 'cp932', 'euc-jp']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # すべて失敗した場合はエラーを無視して読み込み
        return content.decode('utf-8', errors='ignore')
        
    except Exception as e:
        raise ValueError(f"テキストファイルの読み込みに失敗しました: {e}")


def read_uploaded_file(uploaded_file) -> str:
    """
    アップロードされたファイルからテキストを抽出
    
    Args:
        uploaded_file: Streamlitのアップロードファイルオブジェクト
        
    Returns:
        抽出されたテキスト
        
    Raises:
        ValueError: サポートされていないファイル形式の場合
        ImportError: 必要なライブラリがない場合
    """
    filename = uploaded_file.name.lower()
    
    # ファイルをBytesIOに変換（再読み込み可能にする）
    file_bytes = io.BytesIO(uploaded_file.read())
    
    if filename.endswith('.docx'):
        return read_docx(file_bytes)
    elif filename.endswith('.xlsx'):
        return read_xlsx(file_bytes)
    elif filename.endswith('.pdf'):
        return read_pdf(file_bytes)
    elif filename.endswith('.txt'):
        return read_txt(file_bytes)
    else:
        # 拡張子から判定できない場合
        ext = '.' + filename.split('.')[-1] if '.' in filename else '不明'
        supported = ', '.join(SUPPORTED_EXTENSIONS.keys())
        raise ValueError(f"サポートされていないファイル形式です: {ext}\n対応形式: {supported}")


def check_dependencies() -> dict:
    """
    必要なライブラリのインストール状況を確認
    
    Returns:
        ライブラリ名とインストール状況のdict
    """
    dependencies = {
        'python-docx': False,
        'openpyxl': False,
        'pdfplumber': False,
    }
    
    try:
        import docx
        dependencies['python-docx'] = True
    except ImportError:
        pass
    
    try:
        import openpyxl
        dependencies['openpyxl'] = True
    except ImportError:
        pass
    
    try:
        import pdfplumber
        dependencies['pdfplumber'] = True
    except ImportError:
        pass
    
    return dependencies
